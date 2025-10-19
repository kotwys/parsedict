import argparse
import logging
from multiprocessing import Queue, Pool
from pprint import pprint
import sys
from typing import Iterator

from docx import Document
from parsy import ParseError
import yaml

from lexer import Character, extract_characters, extract_entries
import logger
from logger import log
import parser
from output import prepare_for_output


LINE_WIDTH = 72

def report_error(err: ParseError):
    total_length = len(err.stream)
    left = max(
        min(err.index - LINE_WIDTH//2, total_length - LINE_WIDTH),
        0)
    right = min(
        max(left+LINE_WIDTH, err.index + LINE_WIDTH//2),
        total_length)
    text = ''.join(map(lambda c: c.char, err.stream[left:right]))
    log.error('Expected %s\n  %s\n  %s^',
              err.expected, text, ' '*(err.index-left))


def output_raw(result: dict):
    pprint(result, sys.stdout, compact=True)

def output_yaml(result: dict):
    clean = prepare_for_output(result)
    yaml.dump(clean,
              stream=sys.stdout,
              allow_unicode=True,
              sort_keys=False,
              explicit_start=True)


parse_fn = None

def worker_init(
        grammar: str,
        partial_parse: bool,
        log_queue,
        log_level):
    global parse_fn
    import logger
    parser.init_grammar(grammar)
    logger.init_logging(log_queue, log_level)

    if partial_parse:
        parse_fn = lambda x: parser.entry.parse_partial(x)
    else:
        parse_fn = lambda x: parser.entry.parse(x)


def worker_fn(entry: list[Character]):
    logger.entry_prefix = ''.join(map(lambda c: c.char, entry[:10]))
    try:
        result = parse_fn(entry)
        return result
    except ParseError as ex:
        report_error(ex)
        return None


BUFFER_SIZE = 100

def execute_multi_process(jobs, entries, output_fn, *args):
    """Analyses the entries in the specified number of concurrent processes."""
    counter = 0
    buffer = []
    def dump_buffer():
        print('%d entries parsed' % counter, file=sys.stderr)
        for result in buffer:
            output_fn(result)

    with Pool(processes=jobs,
              initializer=worker_init,
              initargs=args) as pool:
        for result in pool.imap_unordered(worker_fn, entries):
            if result:
                counter += 1
                buffer.append(result)
                if len(buffer) >= BUFFER_SIZE:
                    dump_buffer()
                    buffer.clear()

    if buffer:
        dump_buffer()


def execute_single_process(entries, output_fn, *args):
    """Analyses the entries in a single process."""
    worker_init(*args)
    for entry in entries:
        result = worker_fn(entry)
        if result:
            output_fn(result)


def try_parse_int(val: str) -> tuple[int, None] | tuple[None, Exception]:
    try:
        return int(val), None
    except ValueError as ex:
        return None, ex


def main(args: argparse.Namespace):
    if args.partial:
        if args.format != 'raw':
            print('Partial parsing is supported only with raw format.',
                  file=sys.stderr)
            sys.exit(-1)

    output_fn = output_yaml if args.format == 'yaml' else output_raw

    doc = Document(args.infile)
    if args.par is not None:
        idx, err = try_parse_int(args.par)
        if not err:
            # If integer, just find the paragraph by its index
            if idx < 0 or idx >= len(doc.paragraphs):
                print('Paragraph index is out of range (max: %s)'
                      % len(doc.paragraphs)-1,
                      file=sys.stderr)
                sys.exit(-1)
            par = doc.paragraphs[idx]
        else:
            # If string, find the paragraph that begins with that string
            try:
                par = next(p for p in doc.paragraphs
                           if p.text.startswith(args.par))
            except StopIteration:
                print('Paragraph not found', file=sys.stderr)
                sys.exit(-1)
        chars = list(extract_characters(par))
        entries = [chars]
    else:
        entries = extract_entries(doc)

    log_queue = Queue()
    log_level = logging.DEBUG
    listener = logger.init_main_logging(log_queue)

    if args.jobs:
        execute_multi_process(
            args.jobs,
            entries, output_fn,
            args.grammar, args.partial, log_queue, log_level)
    else:
        execute_single_process(
            entries, output_fn,
            args.grammar, args.partial, log_queue, log_level)

    listener.stop()


if __name__ == "__main__":
    p = argparse.ArgumentParser(
        description='Parse formatted text from a .docx file')
    p.add_argument(
        'grammar',
        help='the name of the grammar to use')
    p.add_argument(
        'infile',
        type=argparse.FileType('rb'),
        help='the .docx file')
    p.add_argument(
        '--par',
        help='''parse the given paragraph instead of the full file.
        If PAR is an integer, finds the paragraph with the index PAR.
        Otherwise, finds the first paragraph that begins with PAR.''')
    p.add_argument(
        '--format',
        help='the output format (default is YAML)',
        choices=['raw', 'yaml'],
        default='yaml')
    p.add_argument(
        '--partial',
        action='store_true',
        help='output partial parsing result (only with raw format)')
    p.add_argument(
        '-j', '--jobs',
        type=int,
        default=0,
        help='use this number of processes')
    main(p.parse_args())
